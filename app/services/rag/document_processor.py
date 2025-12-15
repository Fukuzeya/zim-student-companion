# ============================================================================
# Document Processor
# ============================================================================
"""
Sophisticated document processing pipeline with:
- Multiple chunking strategies (semantic, hierarchical, sliding window)
- Intelligent metadata extraction
- Question/answer pair detection for educational content
- Content deduplication
- Quality scoring for chunks
"""
from __future__ import annotations

import hashlib
import re
import unicodedata
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Generator, List, Optional, Tuple
import logging

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.services.rag.config import ChunkingConfig, ChunkingStrategy, get_rag_config

logger = logging.getLogger(__name__)


# ============================================================================
# Data Classes
# ============================================================================
@dataclass
class DocumentChunk:
    """
    Represents a processed chunk of document content.
    Contains content, metadata, and relationships.
    """
    content: str
    metadata: Dict[str, Any]
    chunk_id: str = ""
    parent_chunk_id: Optional[str] = None
    child_chunk_ids: List[str] = field(default_factory=list)
    
    # Quality metrics
    token_count: int = 0
    quality_score: float = 1.0
    
    def __post_init__(self):
        if not self.chunk_id:
            self.chunk_id = self._generate_id()
        if not self.token_count:
            self.token_count = len(self.content.split())
    
    def _generate_id(self) -> str:
        """Generate deterministic chunk ID"""
        content_hash = hashlib.sha256(self.content.encode()).hexdigest()[:12]
        meta_hash = hashlib.md5(str(sorted(self.metadata.items())).encode()).hexdigest()[:8]
        return f"chunk_{content_hash}_{meta_hash}"
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for storage"""
        return {
            "content": self.content,
            "metadata": self.metadata,
            "chunk_id": self.chunk_id,
            "parent_chunk_id": self.parent_chunk_id,
            "child_chunk_ids": self.child_chunk_ids,
            "token_count": self.token_count,
            "quality_score": self.quality_score,
        }


@dataclass
class ZIMSECDocument:
    """
    Represents a ZIMSEC curriculum document with full metadata.
    """
    file_path: str
    document_type: str  # syllabus, past_paper, marking_scheme, textbook, notes
    subject: str
    education_level: str  # primary, secondary, a_level
    grade: str
    
    # Optional metadata
    year: Optional[int] = None
    paper_number: Optional[str] = None  # Paper 1, Paper 2, etc.
    term: Optional[str] = None
    topic: Optional[str] = None
    
    # Processing state
    raw_content: str = ""
    chunks: List[DocumentChunk] = field(default_factory=list)
    processing_metadata: Dict[str, Any] = field(default_factory=dict)
    
    @property
    def source_id(self) -> str:
        """Generate unique source identifier"""
        parts = [self.subject, self.grade, self.document_type]
        if self.year:
            parts.append(str(self.year))
        if self.paper_number:
            parts.append(self.paper_number)
        return "_".join(p.lower().replace(" ", "_") for p in parts)


# ============================================================================
# Chunking Strategies
# ============================================================================
class ChunkingStrategyBase(ABC):
    """Base class for chunking strategies"""
    
    def __init__(self, config: ChunkingConfig):
        self.config = config
    
    @abstractmethod
    def chunk(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Split text into chunks with the strategy"""
        pass
    
    def _clean_chunk(self, text: str) -> str:
        """Clean and normalize chunk text"""
        text = unicodedata.normalize("NFKC", text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    def _estimate_tokens(self, text: str) -> int:
        """Estimate token count (rough approximation)"""
        return len(text.split())


class FixedSizeChunker(ChunkingStrategyBase):
    """Simple fixed-size chunking with overlap"""
    
    def chunk(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        chunks = []
        words = text.split()
        
        i = 0
        chunk_idx = 0
        while i < len(words):
            end = min(i + self.config.chunk_size, len(words))
            chunk_words = words[i:end]
            chunk_text = self._clean_chunk(" ".join(chunk_words))
            
            if len(chunk_words) >= self.config.min_chunk_size:
                chunks.append(DocumentChunk(
                    content=chunk_text,
                    metadata={**metadata, "chunk_index": chunk_idx},
                    token_count=len(chunk_words)
                ))
                chunk_idx += 1
            
            # Move with overlap
            i += self.config.chunk_size - self.config.chunk_overlap
        
        return chunks


class SemanticChunker(ChunkingStrategyBase):
    """Chunk at sentence/paragraph boundaries for semantic coherence"""
    
    # Sentence ending patterns
    SENTENCE_PATTERN = re.compile(r'(?<=[.!?])\s+(?=[A-Z])')
    
    def chunk(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        chunks = []
        
        # Split into paragraphs first
        paragraphs = re.split(r'\n\n+', text)
        
        current_chunk = []
        current_tokens = 0
        chunk_idx = 0
        
        for para in paragraphs:
            para = self._clean_chunk(para)
            if not para:
                continue
            
            para_tokens = self._estimate_tokens(para)
            
            # If single paragraph exceeds max, split by sentences
            if para_tokens > self.config.max_chunk_size:
                sentences = self.SENTENCE_PATTERN.split(para)
                for sent in sentences:
                    sent_tokens = self._estimate_tokens(sent)
                    if current_tokens + sent_tokens > self.config.chunk_size:
                        if current_chunk:
                            chunks.append(self._create_chunk(
                                current_chunk, metadata, chunk_idx
                            ))
                            chunk_idx += 1
                        current_chunk = [sent]
                        current_tokens = sent_tokens
                    else:
                        current_chunk.append(sent)
                        current_tokens += sent_tokens
            elif current_tokens + para_tokens > self.config.chunk_size:
                if current_chunk:
                    chunks.append(self._create_chunk(
                        current_chunk, metadata, chunk_idx
                    ))
                    chunk_idx += 1
                current_chunk = [para]
                current_tokens = para_tokens
            else:
                current_chunk.append(para)
                current_tokens += para_tokens
        
        # Final chunk
        if current_chunk:
            chunks.append(self._create_chunk(current_chunk, metadata, chunk_idx))
        
        return chunks
    
    def _create_chunk(
        self, parts: List[str], metadata: Dict[str, Any], idx: int
    ) -> DocumentChunk:
        content = "\n\n".join(parts)
        return DocumentChunk(
            content=self._clean_chunk(content),
            metadata={**metadata, "chunk_index": idx},
            token_count=self._estimate_tokens(content)
        )


class HierarchicalChunker(ChunkingStrategyBase):
    """
    Creates parent-child chunk relationships.
    Parent chunks provide context, children provide specificity.
    """
    
    def chunk(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        all_chunks = []
        
        # First, create parent chunks (larger context windows)
        parent_chunker = FixedSizeChunker(ChunkingConfig(
            chunk_size=self.config.parent_chunk_size,
            chunk_overlap=100,
            min_chunk_size=200
        ))
        parent_chunks = parent_chunker.chunk(text, metadata)
        
        # For each parent, create child chunks
        for parent_idx, parent in enumerate(parent_chunks):
            parent.metadata["is_parent"] = True
            parent.metadata["hierarchy_level"] = 0
            all_chunks.append(parent)
            
            # Create children from parent content
            child_chunker = SemanticChunker(ChunkingConfig(
                chunk_size=self.config.child_chunk_size,
                chunk_overlap=32,
                min_chunk_size=50
            ))
            children = child_chunker.chunk(parent.content, metadata)
            
            for child in children:
                child.parent_chunk_id = parent.chunk_id
                child.metadata["is_parent"] = False
                child.metadata["hierarchy_level"] = 1
                child.metadata["parent_index"] = parent_idx
                parent.child_chunk_ids.append(child.chunk_id)
                all_chunks.append(child)
        
        return all_chunks


class SlidingWindowChunker(ChunkingStrategyBase):
    """
    Creates overlapping windows with surrounding context.
    Each chunk includes context from neighboring windows.
    """
    
    def chunk(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        chunks = []
        
        # Split into sentences
        sentences = re.split(r'(?<=[.!?])\s+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        if not sentences:
            return []
        
        window_size = self.config.window_size
        stride = self.config.stride
        
        for i in range(0, len(sentences), stride):
            # Main window
            window_end = min(i + window_size, len(sentences))
            window_sentences = sentences[i:window_end]
            
            # Add context (previous sentences as prefix)
            context_start = max(0, i - 2)
            context_sentences = sentences[context_start:i]
            
            # Build chunk with context
            main_content = " ".join(window_sentences)
            
            if context_sentences:
                context = " ".join(context_sentences)
                full_content = f"[Context: {context}]\n\n{main_content}"
            else:
                full_content = main_content
            
            chunk_meta = {
                **metadata,
                "chunk_index": len(chunks),
                "window_start": i,
                "window_end": window_end,
                "has_context": bool(context_sentences)
            }
            
            chunks.append(DocumentChunk(
                content=self._clean_chunk(full_content),
                metadata=chunk_meta,
                token_count=self._estimate_tokens(full_content)
            ))
        
        return chunks


class QuestionAwareChunker(ChunkingStrategyBase):
    """
    Specialized chunker for educational content.
    Detects and preserves question-answer pairs.
    """
    
    # Patterns for detecting questions in educational content
    QUESTION_PATTERNS = [
        r'(?:Question\s*(\d+)[\.\):]?\s*)',
        r'(?:(\d+)[\.\)]\s*(?=[A-Z]))',
        r'(?:Q(\d+)[\.\):]?\s*)',
        r'(?:\(([a-z])\)\s*)',  # Sub-questions like (a), (b)
    ]
    
    def chunk(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        chunks = []
        
        # Try to detect questions
        combined_pattern = '|'.join(self.QUESTION_PATTERNS)
        parts = re.split(combined_pattern, text, flags=re.MULTILINE)
        
        current_question = None
        current_content = []
        question_num = 0
        
        for part in parts:
            if not part or not part.strip():
                continue
            
            # Check if this is a question number
            if re.match(r'^\d+$', part.strip()):
                # Save previous question
                if current_content:
                    chunks.append(self._create_question_chunk(
                        current_content, metadata, current_question, question_num
                    ))
                    question_num += 1
                
                current_question = part.strip()
                current_content = []
            elif re.match(r'^[a-z]$', part.strip()):
                # Sub-question marker - append to current
                current_content.append(f"({part.strip()})")
            else:
                current_content.append(part.strip())
        
        # Handle last question
        if current_content:
            chunks.append(self._create_question_chunk(
                current_content, metadata, current_question, question_num
            ))
        
        # If no questions detected, fall back to semantic chunking
        if not chunks or len(chunks) == 1:
            fallback = SemanticChunker(self.config)
            return fallback.chunk(text, metadata)
        
        return chunks
    
    def _create_question_chunk(
        self,
        parts: List[str],
        metadata: Dict[str, Any],
        question_num: Optional[str],
        idx: int
    ) -> DocumentChunk:
        content = " ".join(parts)
        if question_num:
            content = f"Question {question_num}: {content}"
        
        return DocumentChunk(
            content=self._clean_chunk(content),
            metadata={
                **metadata,
                "chunk_type": "question",
                "question_number": question_num,
                "chunk_index": idx
            },
            token_count=self._estimate_tokens(content)
        )


# ============================================================================
# Content Extractors
# ============================================================================
class PDFExtractor:
    """Extract content from PDF files with structure preservation"""
    
    @staticmethod
    async def extract(file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from PDF.
        
        Returns:
            Tuple of (content, extraction_metadata)
        """
        content_parts = []
        extraction_meta = {
            "page_count": 0,
            "tables_found": 0,
            "images_found": 0,
        }
        
        try:
            with fitz.open(file_path) as pdf:
                extraction_meta["page_count"] = len(pdf)
                
                for page_num, page in enumerate(pdf, 1):
                    # Extract main text
                    text = page.get_text("text")
                    if text.strip():
                        content_parts.append(f"\n[Page {page_num}]\n{text}")
                    
                    # Extract tables
                    try:
                        tables = page.find_tables()
                        for table in tables:
                            extraction_meta["tables_found"] += 1
                            df = table.to_pandas()
                            table_str = df.to_string(index=False)
                            content_parts.append(f"\n[Table - Page {page_num}]\n{table_str}")
                    except Exception:
                        pass  # Table extraction can fail on some PDFs
                    
                    # Count images
                    images = page.get_images()
                    extraction_meta["images_found"] += len(images)
        
        except Exception as e:
            logger.error(f"PDF extraction failed for {file_path}: {e}")
            raise
        
        return "\n".join(content_parts), extraction_meta


class DocxExtractor:
    """Extract content from DOCX files with structure preservation"""
    
    @staticmethod
    async def extract(file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text and metadata from DOCX"""
        content_parts = []
        extraction_meta = {
            "paragraph_count": 0,
            "tables_found": 0,
            "headings_found": 0,
        }
        
        try:
            doc = DocxDocument(file_path)
            
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                
                extraction_meta["paragraph_count"] += 1
                
                # Preserve heading structure
                style_name = para.style.name if para.style else ""
                if style_name.startswith('Heading'):
                    level = style_name.replace('Heading ', '').strip()
                    try:
                        level_int = int(level)
                        content_parts.append(f"\n{'#' * level_int} {para.text}\n")
                        extraction_meta["headings_found"] += 1
                    except ValueError:
                        content_parts.append(para.text)
                else:
                    content_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                extraction_meta["tables_found"] += 1
                content_parts.append("\n[Table]")
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    content_parts.append(" | ".join(cells))
                content_parts.append("[/Table]\n")
        
        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {e}")
            raise
        
        return "\n".join(content_parts), extraction_meta


# ============================================================================
# Main Document Processor
# ============================================================================
class DocumentProcessor:
    """
    Main document processing pipeline.
    Handles extraction, cleaning, chunking, and quality scoring.
    """
    
    def __init__(self, config: Optional[ChunkingConfig] = None):
        self.config = config or get_rag_config().chunking
        
        # Initialize chunking strategies
        self._chunkers: Dict[ChunkingStrategy, ChunkingStrategyBase] = {
            ChunkingStrategy.FIXED_SIZE: FixedSizeChunker(self.config),
            ChunkingStrategy.SEMANTIC: SemanticChunker(self.config),
            ChunkingStrategy.HIERARCHICAL: HierarchicalChunker(self.config),
            ChunkingStrategy.SLIDING_WINDOW: SlidingWindowChunker(self.config),
            ChunkingStrategy.QUESTION_AWARE: QuestionAwareChunker(self.config),
        }
    
    async def process_document(
        self,
        doc: ZIMSECDocument,
        strategy: Optional[ChunkingStrategy] = None
    ) -> ZIMSECDocument:
        """
        Process a document through the full pipeline.
        
        Args:
            doc: Document to process
            strategy: Override chunking strategy (auto-selects if None)
        
        Returns:
            Processed document with chunks
        """
        start_time = datetime.now()
        file_path = Path(doc.file_path)
        
        logger.info(f"Processing document: {file_path.name}")
        
        # Step 1: Extract content
        doc.raw_content, extraction_meta = await self._extract_content(file_path)
        doc.processing_metadata["extraction"] = extraction_meta
        
        # Step 2: Clean and preprocess
        doc.raw_content = self._preprocess_content(doc.raw_content)
        
        # Step 3: Select chunking strategy
        chunking_strategy = strategy or self._select_strategy(doc)
        logger.info(f"Using chunking strategy: {chunking_strategy.value}")
        
        # Step 4: Create chunks
        base_metadata = self._build_base_metadata(doc)
        chunker = self._chunkers[chunking_strategy]
        doc.chunks = chunker.chunk(doc.raw_content, base_metadata)
        
        # Step 5: Score chunk quality
        doc.chunks = self._score_chunks(doc.chunks)
        
        # Step 6: Deduplicate
        doc.chunks = self._deduplicate_chunks(doc.chunks)
        
        # Record processing metadata
        doc.processing_metadata.update({
            "chunking_strategy": chunking_strategy.value,
            "total_chunks": len(doc.chunks),
            "processing_time_ms": (datetime.now() - start_time).total_seconds() * 1000,
            "avg_chunk_tokens": sum(c.token_count for c in doc.chunks) / len(doc.chunks) if doc.chunks else 0,
        })
        
        logger.info(
            f"Processed {file_path.name}: {len(doc.chunks)} chunks, "
            f"{doc.processing_metadata['processing_time_ms']:.0f}ms"
        )
        
        return doc
    
    async def _extract_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract content based on file type"""
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return await PDFExtractor.extract(file_path)
        elif suffix == '.docx':
            return await DocxExtractor.extract(file_path)
        elif suffix == '.txt':
            content = file_path.read_text(encoding='utf-8')
            return content, {"file_type": "text"}
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    
    def _preprocess_content(self, content: str) -> str:
        """Clean and normalize document content"""
        # Normalize unicode
        content = unicodedata.normalize("NFKC", content)
        
        # Remove excessive whitespace
        content = re.sub(r'\n{4,}', '\n\n\n', content)
        content = re.sub(r' {3,}', '  ', content)
        content = re.sub(r'\t+', ' ', content)
        
        # Normalize bullet points
        content = re.sub(r'^[\s]*[•●○▪►◦‣⁃]\s*', '- ', content, flags=re.MULTILINE)
        
        # Clean common OCR artifacts
        content = re.sub(r'["""]', '"', content)
        content = re.sub(r"[''']", "'", content)
        content = re.sub(r'…', '...', content)
        
        # Remove page numbers (common patterns)
        content = re.sub(r'\n\s*-?\s*\d+\s*-?\s*\n', '\n', content)
        
        return content.strip()
    
    def _select_strategy(self, doc: ZIMSECDocument) -> ChunkingStrategy:
        """Auto-select best chunking strategy based on document type"""
        strategy_map = {
            "past_paper": ChunkingStrategy.QUESTION_AWARE,
            "marking_scheme": ChunkingStrategy.QUESTION_AWARE,
            "syllabus": ChunkingStrategy.HIERARCHICAL,
            "textbook": ChunkingStrategy.HIERARCHICAL,
            "teacher_notes": ChunkingStrategy.SEMANTIC,
        }
        return strategy_map.get(doc.document_type, self.config.strategy)
    
    def _build_base_metadata(self, doc: ZIMSECDocument) -> Dict[str, Any]:
        """Build base metadata for all chunks"""
        return {
            "source_file": doc.file_path,
            "source_id": doc.source_id,
            "document_type": doc.document_type,
            "subject": doc.subject,
            "education_level": doc.education_level,
            "grade": doc.grade,
            "year": doc.year,
            "paper_number": doc.paper_number,
            "term": doc.term,
            "topic": doc.topic,
            "processed_at": datetime.now().isoformat(),
        }
    
    def _score_chunks(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """Score chunk quality based on multiple factors"""
        for chunk in chunks:
            score = 1.0
            
            # Penalize very short chunks
            if chunk.token_count < 50:
                score *= 0.7
            
            # Penalize chunks that are mostly whitespace or punctuation
            content = chunk.content
            alpha_ratio = sum(c.isalpha() for c in content) / max(len(content), 1)
            if alpha_ratio < 0.5:
                score *= 0.8
            
            # Boost chunks with educational keywords
            edu_keywords = ['explain', 'define', 'calculate', 'describe', 'compare',
                          'example', 'formula', 'method', 'step', 'answer']
            keyword_count = sum(1 for kw in edu_keywords if kw in content.lower())
            score *= (1.0 + keyword_count * 0.05)
            
            # Boost question-containing chunks
            if '?' in content or chunk.metadata.get('chunk_type') == 'question':
                score *= 1.1
            
            chunk.quality_score = min(score, 2.0)  # Cap at 2.0
        
        return chunks
    
    def _deduplicate_chunks(
        self,
        chunks: List[DocumentChunk],
        similarity_threshold: float = 0.95
    ) -> List[DocumentChunk]:
        """Remove near-duplicate chunks based on content similarity"""
        if len(chunks) <= 1:
            return chunks
        
        unique_chunks = []
        seen_hashes = set()
        
        for chunk in chunks:
            # Create a normalized hash for comparison
            normalized = re.sub(r'\s+', ' ', chunk.content.lower().strip())
            content_hash = hashlib.md5(normalized.encode()).hexdigest()
            
            if content_hash not in seen_hashes:
                seen_hashes.add(content_hash)
                unique_chunks.append(chunk)
        
        removed = len(chunks) - len(unique_chunks)
        if removed > 0:
            logger.info(f"Deduplicated {removed} chunks")
        
        return unique_chunks
    
    # ==================== Batch Processing ====================
    
    async def process_directory(
        self,
        directory: Path,
        document_type: str,
        subject: Optional[str] = None,
        education_level: str = "secondary",
        **kwargs
    ) -> List[ZIMSECDocument]:
        """Process all documents in a directory"""
        processed = []
        
        for file_path in directory.glob("**/*"):
            if file_path.suffix.lower() in ['.pdf', '.docx', '.txt']:
                try:
                    # Parse metadata from filename if not provided
                    meta = self._parse_filename(file_path.name)
                    
                    doc = ZIMSECDocument(
                        file_path=str(file_path),
                        document_type=document_type,
                        subject=subject or meta.get("subject", "unknown"),
                        education_level=education_level,
                        grade=meta.get("grade", "Form 3"),
                        year=meta.get("year"),
                        paper_number=meta.get("paper"),
                        **kwargs
                    )
                    
                    processed_doc = await self.process_document(doc)
                    processed.append(processed_doc)
                    
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
        
        return processed
    
    def _parse_filename(self, filename: str) -> Dict[str, Any]:
        """Extract metadata from filename conventions"""
        # Expected: subject_level_grade_year_paper.pdf
        parts = filename.lower().replace('.pdf', '').replace('.docx', '').replace('.txt', '').split('_')
        
        result = {}
        if len(parts) >= 1:
            result["subject"] = parts[0].replace("-", " ").title()
        if len(parts) >= 3:
            result["grade"] = parts[2].replace("-", " ").title()
        if len(parts) >= 4 and parts[3].isdigit():
            result["year"] = int(parts[3])
        if len(parts) >= 5:
            result["paper"] = parts[4].replace("-", " ").title()
        
        return result